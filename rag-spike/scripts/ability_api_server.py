from __future__ import annotations

import argparse
import io
import json
import textwrap
import time
import warnings
import zipfile
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import extract_role_profile

warnings.filterwarnings("ignore", message="'cgi' is deprecated.*", category=DeprecationWarning)
import cgi  # noqa: E402


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_HOST = "127.0.0.1"
DEFAULT_PORT = 8766

CAPABILITIES = [
    {
        "key": "communication_expression",
        "label": "沟通表达能力",
        "description": "清晰表达、结构化讲述、回应问题、传递信息。",
    },
    {
        "key": "logical_analysis",
        "label": "逻辑分析能力",
        "description": "拆解问题、识别因果、建立框架、形成判断。",
    },
    {
        "key": "learning_adaptability",
        "label": "学习适应能力",
        "description": "快速学习新知识、新工具、新环境，并调整方法。",
    },
    {
        "key": "execution_ownership",
        "label": "执行推进能力",
        "description": "目标拆解、行动推进、按时交付、承担结果。",
    },
    {
        "key": "collaboration_leadership",
        "label": "协作与领导力",
        "description": "团队协作、协调资源、推动他人、处理分歧。",
    },
    {
        "key": "self_awareness_motivation",
        "label": "自我认知与职业动机",
        "description": "理解自身优势、短板、兴趣、价值观和职业选择原因。",
    },
    {
        "key": "data_digital_literacy",
        "label": "数据与数字化思维",
        "description": "使用数据、工具、数字化方法支持分析和决策。",
    },
    {
        "key": "business_industry_understanding",
        "label": "商业与行业理解",
        "description": "理解行业、商业模式、用户需求、组织目标和市场环境。",
    },
]

CAPABILITY_KEY_LIST = [item["key"] for item in CAPABILITIES]
CAPABILITY_KEYS = set(CAPABILITY_KEY_LIST)
SUPPORTED_RESUME_TYPES = {".docx", ".pdf"}


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as error:
        return extract_docx_text_with_stdlib(data, error)

    document = Document(io.BytesIO(data))
    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_docx_text_with_stdlib(data: bytes, import_error: ImportError) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            document_xml = archive.read("word/document.xml")
    except Exception as error:  # noqa: BLE001 - convert archive/XML issues to user-facing API error
        raise RuntimeError("缺少 python-docx 依赖，且无法用标准库解析该 .docx。") from import_error or error

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = ElementTree.fromstring(document_xml)
    paragraphs: list[str] = []
    for paragraph in root.findall(".//w:p", namespace):
        texts = [node.text or "" for node in paragraph.findall(".//w:t", namespace)]
        line = "".join(texts).strip()
        if line:
            paragraphs.append(line)
    return "\n".join(paragraphs).strip()


def extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise RuntimeError("缺少 pypdf 依赖，请先安装 requirements.txt。") from error

    reader = PdfReader(io.BytesIO(data))
    parts = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n".join(part for part in parts if part).strip()
    if not text:
        raise ValueError("该 PDF 可能是扫描件，请粘贴文字版简历。")
    return text


def extract_resume_text(file_name: str, data: bytes) -> dict[str, Any]:
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_RESUME_TYPES:
        raise ValueError("仅支持 .docx 和文字版 .pdf 简历文件。")
    if not data:
        raise ValueError("上传文件为空，请重新选择文件。")

    if suffix == ".docx":
        text = extract_docx_text(data)
    else:
        text = extract_pdf_text(data)

    if not text:
        raise ValueError("未能从文件中提取文字，请粘贴文字版简历。")

    return {
        "file_name": file_name,
        "file_type": suffix.lstrip("."),
        "text": text,
        "extraction_status": "extracted",
    }


def build_capability_prompt(
    *,
    user_id: str,
    resume_text: str,
    target_role: str,
    questionnaire_answers: list[dict[str, Any]],
    role_requirements: dict[str, Any] | None = None,
    role_dimensions: list[dict[str, Any]] | None = None,
) -> str:
    capability_lines = "\n".join(
        f"- {item['key']}: {item['label']}。{item['description']}" for item in CAPABILITIES
    )
    questionnaire_lines = "\n".join(
        f"- {item.get('id', f'qa_{index + 1:02d}')} | "
        f"{item.get('capability_key', '')} | {item.get('indicator', '')} | "
        f"score={item.get('score', 'not_provided')} | reverse={item.get('reverse', False)}\n"
        f"  题目：{item.get('text', '')}"
        for index, item in enumerate(questionnaire_answers)
    )
    role_requirement_lines = "\n".join(
        f"- {key}: required_level={value.get('required_level', 'unknown')}, "
        f"weight={value.get('weight', 'unknown')}, "
        f"summary={value.get('requirement_summary', '')}"
        for key, value in (role_requirements or {}).items()
        if key in CAPABILITY_KEYS and isinstance(value, dict)
    )
    role_dimension_lines = "\n".join(
        "- {dimension_id} | {label} | required_level={required_level} | weight={weight} | mapped={mapped}\n"
        "  evaluation: {evaluation}\n"
        "  focus: {focus}".format(
            dimension_id=item.get("dimension_id", ""),
            label=item.get("label", ""),
            required_level=item.get("required_level", ""),
            weight=item.get("weight", ""),
            mapped=", ".join(str(key) for key in item.get("mapped_capability_keys", [])),
            evaluation=item.get("evaluation_method", ""),
            focus=item.get("questionnaire_focus", ""),
        )
        for item in role_dimensions or []
        if isinstance(item, dict)
    )

    return textwrap.dedent(
        f"""
        你是学生职业能力画像评分器。请根据简历文本、心仪职业和 48 题行为锚定自评问卷，
        输出可合并进 capability_profile 的局部能力证据 JSON，并为每个能力给出可马上执行的改进建议。

        只允许输出 JSON，不要输出 Markdown，不要解释。
        不允许输出岗位推荐列表、录用判断、医学诊断、心理测评结论或简历正文改写；除 report_content.improvement_plan 外不要输出额外训练计划。
        如果证据不足，可以给较低 confidence；不要把证据不足直接写成能力差。
        自评问卷只能作为低可信度辅助证据，不能单独决定最终能力水平。
        简历文本和自评分数必须分开解释。

        用户：
        user_id: {user_id}
        target_role: {target_role or "未提供"}

        只能使用以下 capability_key：
        {capability_lines}

        岗位能力要求（如提供，必须用于生成 improvement_advice）：
        {role_requirement_lines or "未提供岗位能力要求，仅根据 target_role、简历和问卷生成建议。"}

        岗位专属能力维度（如提供，用于理解当前岗位的 6 维模型；输出仍必须使用 capability_key）：
        {role_dimension_lines or "未提供岗位专属能力维度。"}

        输出要求：
        - 顶层必须包含 evidence 数组。
        - 顶层必须包含 report_content 对象，用于前端直接展示报告文案。
        - evidence 数组对象的 source_type 只能是 resume_text 或 self_assessment。
        - 每个 evidence 对象必须包含 source_type、source_id、capability_evidence。
        - capability_evidence 中每项必须包含 capability_key、score、confidence、evidence_summary、improvement_advice。
        - score 范围 0-100。
        - confidence 范围 0.00-1.00。
        - evidence_summary 必须说明评分依据；证据不足时必须明确写出缺少什么行为证据。
        - improvement_advice 必须是中文，必须针对 target_role、该 capability_key、简历证据和岗位能力要求。
        - improvement_advice 必须落到今天或本周可以做的具体动作；不要写“加强学习”“提升能力”“多练习”这类空话。
        - improvement_advice 不要超过 80 个中文字符，优先建议补充经历证据、量化材料、面试表达稿、小案例或具体工具练习。
        - resume_text 证据要基于简历里的项目、行为、结果、数据或职责。
        - self_assessment 证据要明确写成“自评倾向”，confidence 通常不高于 0.45。
        - 当自评高分但简历缺少对应行为证据时，要提示“需要补充行为证据”。
        - report_content.capability_details 必须围绕岗位专属能力维度生成；如提供 role_dimensions，必须尽量一一对应 role_dimension_id。
        - report_content.capability_details 每项必须包含 role_dimension_id、role_application、personal_assessment、improvement_advice。
        - role_application、personal_assessment、improvement_advice 必须是完整中文自然段，不能使用前端模板式开头。
        - report_content.improvement_plan 必须是 4 个模块，围绕 4 周执行计划生成；模块可参考工具提升、基础能力、行业知识、实践项目，但具体内容必须结合 target_role、简历证据、问卷答案和岗位差距生成。
        - report_content.improvement_plan 每项必须包含 title 和 items；items 每条必须是可执行动作。
        - report_content 中所有展示文案可以使用 Markdown 粗体标记 **重点内容**，由你根据内容判断重点。
        - 粗体只标真正关键的对象、差距、动作或交付物；不要把固定句式、整段文字或标签机械加粗。
        - 不要输出 HTML，不要输出 Markdown 标题，只允许在字符串内部使用 **...** 表示加粗。

        目标 JSON 形状：
        {{
          "evidence": [
            {{
              "source_type": "resume_text",
              "source_id": "resume_text_01",
              "capability_evidence": [
                {{
                  "capability_key": "communication_expression",
                  "score": 72,
                  "confidence": 0.58,
                  "evidence_summary": "简历显示用户能描述项目背景和个人行动，但结果指标证据不足。",
                  "improvement_advice": "今天选一个项目写 90 秒表达稿，补齐背景、个人动作、结果指标和复盘。"
                }}
              ]
            }},
            {{
              "source_type": "self_assessment",
              "source_id": "questionnaire_48",
              "capability_evidence": [
                {{
                  "capability_key": "communication_expression",
                  "score": 78,
                  "confidence": 0.36,
                  "evidence_summary": "问卷体现较高自评倾向，但仍需要简历或行为案例支持。",
                  "improvement_advice": "本周补一个真实沟通案例，写清楚沟通对象、分歧点、你的表达和结果。"
                }}
              ]
            }}
          ],
          "report_content": {{
            "capability_details": [
              {{
                "role_dimension_id": "requirement_analysis",
                "role_application": "该岗位最看重能否把模糊需求转成**清晰的用户场景和优先级**，并让设计、研发能据此推进。",
                "personal_assessment": "目前证据显示你有技术实现经历，但缺少**用户需求拆解和优先级判断**案例，因此这一项更像是待补强能力。",
                "improvement_advice": "本周选择一个熟悉功能，补一页**用户痛点、约束、优先级和验收标准**，形成可放入作品集的需求拆解页。"
              }}
            ],
            "improvement_plan": [
              {{
                "title": "工具提升（第 1 周）",
                "items": ["每天 1 小时完成一个和目标岗位相关的工具产出，重点沉淀**可展示作品**。"]
              }},
              {{
                "title": "基础能力提升（第 2 周）",
                "items": ["围绕最弱能力做 3 次拆解训练，每次写出**问题、判断依据和复盘结论**。"]
              }},
              {{
                "title": "行业基础知识（贯穿全程）",
                "items": ["每天 30 分钟阅读岗位相关资料，整理**术语、指标和典型业务场景**。"]
              }},
              {{
                "title": "实践项目（第 3-4 周）",
                "items": ["完成一个小项目并输出**原型、分析文档或作品集案例**，最后准备 90 秒讲解。"]
              }}
            ]
          }}
        }}

        简历文本：
        {resume_text or "未提供"}

        问卷答案：
        {questionnaire_lines or "未提供有效问卷答案"}
        """
    ).strip()


def validate_capability_evidence(payload: dict[str, Any], *, require_report_content: bool = False) -> list[str]:
    errors: list[str] = []
    evidence = payload.get("evidence")
    if not isinstance(evidence, list) or not evidence:
        return ["evidence must be a non-empty array"]

    for group_index, group in enumerate(evidence):
        if not isinstance(group, dict):
            errors.append(f"evidence[{group_index}] must be an object")
            continue

        source_type = group.get("source_type")
        source_id = group.get("source_id")
        items = group.get("capability_evidence")
        if source_type not in {"resume_text", "self_assessment", "questionnaire"}:
            errors.append(f"Invalid source_type at evidence[{group_index}]")
        if not isinstance(source_id, str) or not source_id.strip():
            errors.append(f"Missing source_id at evidence[{group_index}]")
        if not isinstance(items, list) or not items:
            errors.append(f"capability_evidence must be non-empty at evidence[{group_index}]")
            continue

        for item_index, item in enumerate(items):
            if not isinstance(item, dict):
                errors.append(
                    f"capability_evidence[{item_index}] at evidence[{group_index}] must be an object"
                )
                continue
            key = item.get("capability_key")
            score = item.get("score")
            confidence = item.get("confidence")
            summary = item.get("evidence_summary")
            advice = item.get("improvement_advice")
            if key not in CAPABILITY_KEYS:
                errors.append(f"Unknown capability_key: {key}")
            if not isinstance(score, (int, float)) or not 0 <= float(score) <= 100:
                errors.append(f"Invalid score for {key}")
            if not isinstance(confidence, (int, float)) or not 0 <= float(confidence) <= 1:
                errors.append(f"Invalid confidence for {key}")
            if not isinstance(summary, str) or not summary.strip():
                errors.append(f"Missing evidence_summary for {key}")
            if not isinstance(advice, str) or not advice.strip():
                errors.append(f"Missing improvement_advice for {key}")

    report_content = payload.get("report_content")
    if require_report_content or report_content is not None:
        if not isinstance(report_content, dict):
            errors.append("report_content must be an object")
            return errors

        details = report_content.get("capability_details")
        if not isinstance(details, list) or not details:
            errors.append("report_content.capability_details must be a non-empty array")
        else:
            for index, item in enumerate(details):
                if not isinstance(item, dict):
                    errors.append(f"report_content.capability_details[{index}] must be an object")
                    continue
                for field in ("role_dimension_id", "role_application", "personal_assessment", "improvement_advice"):
                    value = item.get(field)
                    if not isinstance(value, str) or not value.strip():
                        errors.append(f"Missing {field} at report_content.capability_details[{index}]")

        plan = report_content.get("improvement_plan")
        if not isinstance(plan, list) or len(plan) != 4:
            errors.append("report_content.improvement_plan must contain exactly 4 sections")
        else:
            for index, section in enumerate(plan):
                if not isinstance(section, dict):
                    errors.append(f"report_content.improvement_plan[{index}] must be an object")
                    continue
                title = section.get("title")
                items = section.get("items")
                if not isinstance(title, str) or not title.strip():
                    errors.append(f"Missing title at report_content.improvement_plan[{index}]")
                if not isinstance(items, list) or not items:
                    errors.append(f"Missing items at report_content.improvement_plan[{index}]")
                elif not all(isinstance(item, str) and item.strip() for item in items):
                    errors.append(f"Invalid items at report_content.improvement_plan[{index}]")
    return errors


def score_capability_for_request(
    *,
    user_id: str,
    resume_text: str,
    target_role: str,
    questionnaire_answers: list[dict[str, Any]],
    timeout: int,
    retries: int,
    role_requirements: dict[str, Any] | None = None,
    role_dimensions: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    started_at = time.perf_counter()
    config = extract_role_profile.get_config(extract_role_profile.ENV_PATH)
    prompt = build_capability_prompt(
        user_id=user_id,
        resume_text=resume_text,
        target_role=target_role,
        questionnaire_answers=questionnaire_answers,
        role_requirements=role_requirements,
        role_dimensions=role_dimensions,
    )

    last_error: Exception | None = None
    parsed: dict[str, Any] | None = None
    for attempt in range(retries + 1):
        try:
            content = extract_role_profile.call_deepseek(config, prompt, timeout)
            parsed = extract_role_profile.extract_json_response(content)
            validation_errors = validate_capability_evidence(parsed, require_report_content=True)
            if validation_errors:
                raise RuntimeError("; ".join(validation_errors))
            break
        except Exception as error:  # noqa: BLE001 - local demo API returns structured errors
            last_error = error
            if attempt >= retries:
                raise RuntimeError(f"DeepSeek capability scoring failed: {error}") from error
            time.sleep(1)

    if parsed is None:
        raise RuntimeError(f"DeepSeek capability scoring failed: {last_error}")

    return {
        "evidence": parsed["evidence"],
        "report_content": parsed.get("report_content") or {},
        "deepseek_model": config["model"],
        "validation_errors": validate_capability_evidence(parsed, require_report_content=True),
        "elapsed_seconds": round(time.perf_counter() - started_at, 3),
        "llm_status": "llm_generated_capability_evidence",
    }


class AbilityApiHandler(BaseHTTPRequestHandler):
    server_version = "CapabilityAbilityApi/0.2"

    def do_OPTIONS(self) -> None:  # noqa: N802 - BaseHTTPRequestHandler hook
        self.send_json({}, status=204)

    def do_GET(self) -> None:  # noqa: N802 - BaseHTTPRequestHandler hook
        if self.path.split("?", 1)[0] != "/health":
            self.send_json({"error": "Not found"}, status=404)
            return
        self.send_json(
            {
                "ok": True,
                "service": "capability-assessment-ability-api",
                "endpoints": ["/capability-evidence", "/resume-text"],
            }
        )

    def do_POST(self) -> None:  # noqa: N802 - BaseHTTPRequestHandler hook
        path = self.path.split("?", 1)[0]
        if path == "/resume-text":
            self.handle_resume_text()
            return
        if path == "/capability-evidence":
            self.handle_capability_evidence()
            return
        self.send_json({"error": "Not found"}, status=404)

    def handle_resume_text(self) -> None:
        try:
            content_type = self.headers.get("Content-Type", "")
            if not content_type.lower().startswith("multipart/form-data"):
                raise ValueError("Content-Type must be multipart/form-data.")

            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": content_type,
                },
            )
            file_item = form["file"] if "file" in form else None
            if isinstance(file_item, list):
                file_item = file_item[0] if file_item else None
            if file_item is None or not getattr(file_item, "filename", ""):
                raise ValueError("请上传 .docx 或文字版 .pdf 简历文件。")

            file_name = Path(file_item.filename).name
            data = file_item.file.read()
            self.send_json(extract_resume_text(file_name, data))
        except Exception as error:  # noqa: BLE001 - return structured local demo errors
            self.send_json(
                {
                    "error": str(error),
                    "text": "",
                    "extraction_status": "failed",
                },
                status=400,
            )

    def handle_capability_evidence(self) -> None:
        try:
            payload = self.read_json_payload()
            user_id = str(payload.get("user_id") or "demo_user_001").strip()
            resume_text = str(payload.get("resume_text") or "").strip()
            target_role = str(payload.get("target_role") or "").strip()
            questionnaire_answers = payload.get("questionnaire_answers") or []
            role_requirements = payload.get("role_requirements") or None
            role_dimensions = payload.get("role_dimensions") or None
            timeout = int(payload.get("timeout") or 90)
            retries = int(payload.get("retries") or 1)

            if not resume_text:
                raise ValueError("resume_text is required.")
            if not target_role:
                raise ValueError("target_role is required.")
            if not isinstance(questionnaire_answers, list) or not questionnaire_answers:
                raise ValueError("questionnaire_answers must be a non-empty array.")
            if not 5 <= timeout <= 180:
                raise ValueError("timeout must be between 5 and 180.")
            if not 0 <= retries <= 3:
                raise ValueError("retries must be between 0 and 3.")

            result = score_capability_for_request(
                user_id=user_id,
                resume_text=resume_text,
                target_role=target_role,
                questionnaire_answers=questionnaire_answers,
                timeout=timeout,
                retries=retries,
                role_requirements=role_requirements if isinstance(role_requirements, dict) else None,
                role_dimensions=role_dimensions if isinstance(role_dimensions, list) else None,
            )
            self.send_json(result)
        except Exception as error:  # noqa: BLE001 - return structured local demo errors
            self.send_json(
                {
                    "error": str(error),
                    "evidence": [],
                    "validation_errors": [str(error)],
                    "llm_status": "error",
                },
                status=500,
            )

    def read_json_payload(self) -> dict[str, Any]:
        content_length = int(self.headers.get("Content-Length") or 0)
        body = self.rfile.read(content_length).decode("utf-8")
        if not body:
            return {}
        payload = json.loads(body)
        if not isinstance(payload, dict):
            raise ValueError("JSON body must be an object.")
        return payload

    def send_json(self, payload: dict[str, Any], status: int = 200) -> None:
        body = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        if status != 204:
            self.wfile.write(body)

    def log_message(self, format: str, *args: Any) -> None:
        print(f"[ability-api] {self.address_string()} - {format % args}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Run local ability scoring API server.")
    parser.add_argument("--host", default=DEFAULT_HOST)
    parser.add_argument("--port", type=int, default=DEFAULT_PORT)
    args = parser.parse_args()

    server = ThreadingHTTPServer((args.host, args.port), AbilityApiHandler)
    print(f"Ability API server listening on http://{args.host}:{args.port}")
    print("Open /health to verify the server. Press Ctrl+C to stop.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("Stopping ability API server.")
    finally:
        server.server_close()


if __name__ == "__main__":
    main()
